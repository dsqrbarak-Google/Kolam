import os
import json
import time
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

app = FastAPI(title="קולם - שרת מקומי")

# Enable CORS for local testing
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

CONFIG_FILE = "config.json"

def cleanup_old_backups(backup_dir):
    try:
        if not os.path.exists(backup_dir):
            return
        now = time.time()
        one_week_ago = now - (7 * 24 * 3600)
        for filename in os.listdir(backup_dir):
            if filename.endswith(".wav"):
                filepath = os.path.join(backup_dir, filename)
                if os.path.isfile(filepath):
                    mtime = os.path.getmtime(filepath)
                    if mtime < one_week_ago:
                        os.remove(filepath)
                        print(f"Deleted backup WAV older than 7 days: {filepath}")
    except Exception as e:
        print(f"Failed to cleanup old backup files: {e}")

def load_config():
    if not os.path.exists(CONFIG_FILE):
        config_data = {
            "gemini_api_key": "",
            "document_path": "C:/Kolam/my_biography.docx",
            "backup_audio_dir": "C:/Kolam/backup_audio",
            "user_name": "אבא",
            "version": "0.0.13"
        }
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(config_data, f, indent=2, ensure_ascii=False)
        return config_data
        
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        config_data = json.load(f)
        
    # Trigger cleanup
    cleanup_old_backups(config_data.get("backup_audio_dir", "C:/Kolam/backup_audio"))
    return config_data

# Load config on startup
config = load_config()

# Ensure backup directory exists
backup_dir = config.get("backup_audio_dir", "C:/Stories/backup_audio")
os.makedirs(backup_dir, exist_ok=True)

@app.get("/api/config")
def get_config():
    # Reload config to get latest updates
    global config
    config = load_config()
    doc_path = config.get("document_path", "C:/Kolam/my_biography.docx")
    story_dir = os.path.dirname(doc_path)
    active_story = os.path.basename(doc_path)
    return {
        "user_name": config.get("user_name", "אבא"),
        "version": config.get("version", "0.0.13"),
        "document_path": doc_path,
        "story_dir": story_dir,
        "active_story": active_story,
        "gemini_api_key": config.get("gemini_api_key", ""),
        "has_api_key": bool(config.get("gemini_api_key") and config.get("gemini_api_key") != "YOUR_GEMINI_API_KEY")
    }

@app.post("/api/upload-audio")
async def upload_audio(file: UploadFile = File(...)):
    global config
    config = load_config()
    
    # Generate filename with timestamp
    timestamp = int(time.time())
    backup_filename = f"recording_{timestamp}.wav"
    backup_path = os.path.join(backup_dir, backup_filename)
    
    # Read file content
    audio_bytes = await file.read()
    
    # Save the file locally first (serves as the primary backup and file log)
    try:
        with open(backup_path, "wb") as f:
            f.write(audio_bytes)
        print(f"Saved audio log to {backup_path}")
    except Exception as e:
        print(f"Failed to save audio file to disk: {e}")
        # Continue anyway to try to transcribe

    api_key = config.get("gemini_api_key", "").strip()
    if not api_key or api_key == "YOUR_GEMINI_API_KEY":
        return JSONResponse(
            status_code=400,
            content={
                "status": "error",
                "message": "מפתח ה-API של Gemini לא הוגדר. אנא הגדר אותו בקובץ config.json",
                "saved_backup": True,
                "backup_path": backup_path
            }
        )

    # Call Gemini API
    try:
        client = genai.Client(api_key=api_key)
        
        # Audio transcription prompt
        prompt = (
            "אתה עוזר כתיבה מקצועי לביוגרפיות. לפניך קובץ שמע של אדם מבוגר המכתיב את סיפור חייו בעברית. "
            "אנא תמלל את השמע במדויק. הפוך את הדיבור לטקסט ספרותי, קולח ומרגש, אך הקפד לשמור על הקול האישי, "
            "הסגנון, ביטויי הדיבור הייחודיים ורוח הדברים של הדובר. תקן שגיאות דקדוק קלות או חזרות מיותרות "
            "במידת הצורך כדי לשפר את זרימת הקריאה, אך אל תשנה את המשמעות ואל תוסיף שום הערה או פרשנות משלך. "
            "פלוט אך ורק את הטקסט המתומלל והערוך ללא שום תוספת."
        )
        
        print("Sending audio to Gemini for transcription...")
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                types.Part.from_bytes(
                    data=audio_bytes,
                    mime_type='audio/wav',
                ),
                prompt
            ]
        )
        
        transcript_text = response.text.strip() if response.text else ""
        if not transcript_text:
            raise ValueError("תמלול ריק התקבל מג'מיני")
            
        print(f"Transcription successful: {transcript_text[:100]}...")
        
        # Append to DOCX file
        doc_path = config.get("document_path", "C:/Stories/my_biography.docx")
        os.makedirs(os.path.dirname(doc_path), exist_ok=True)
        
        if os.path.exists(doc_path):
            doc = docx.Document(doc_path)
        else:
            doc = docx.Document()
            title_p = doc.add_paragraph()
            title_p.paragraph_format.rtl = True
            title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = title_p.add_run(f"הביוגרפיה של {config.get('user_name', 'סבא')}")
            run.bold = True
            run.font.size = docx.shared.Pt(24)
            doc.add_paragraph() # Add empty line
            
        # Add new paragraph
        p = doc.add_paragraph()
        p.paragraph_format.rtl = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_run = p.add_run(transcript_text)
        p_run.font.size = docx.shared.Pt(14)
        
        doc.save(doc_path)
        print(f"Appended paragraph to document: {doc_path}")
        
        return {
            "status": "success",
            "transcript": transcript_text,
            "document_path": doc_path
        }
        
    except Exception as e:
        print(f"Error during Gemini API call or DOCX write: {e}")
        # Fail-safe already saved the WAV as backup_path
        return JSONResponse(
            status_code=500,
            content={
                "status": "error",
                "message": f"שגיאת עיבוד או חיבור: {str(e)}",
                "saved_backup": True,
                "backup_path": backup_path
            }
        )

@app.get("/api/last-paragraphs")
def get_last_paragraphs():
    global config
    config = load_config()
    doc_path = config.get("document_path", "C:/Stories/my_biography.docx")
    
    if not os.path.exists(doc_path):
        return {"status": "empty", "paragraphs": []}
        
    try:
        doc = docx.Document(doc_path)
        # Filter out empty paragraphs and the title paragraph if it contains 'הביוגרפיה של'
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text and not text.startswith("הביוגרפיה של"):
                paragraphs.append(text)
                
        # Get last 3 paragraphs
        last_3 = paragraphs[-3:] if len(paragraphs) >= 3 else paragraphs
        return {"status": "success", "paragraphs": last_3}
    except Exception as e:
        print(f"Error reading docx: {e}")
        return {"status": "error", "message": str(e), "paragraphs": []}

@app.post("/api/update-config")
async def update_config(data: dict):
    global config
    config = load_config()
    
    if "user_name" in data:
        config["user_name"] = data["user_name"]
    if "gemini_api_key" in data:
        config["gemini_api_key"] = data["gemini_api_key"]
        
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, ensure_ascii=False)
        
    # Reload directories in case paths changed
    global backup_dir
    backup_dir = config.get("backup_audio_dir", "C:/Stories/backup_audio")
    os.makedirs(backup_dir, exist_ok=True)
    
    doc_path = config.get("document_path", "C:/Kolam/my_biography.docx")
    story_dir = os.path.dirname(doc_path)
    active_story = os.path.basename(doc_path)
    
    return {"status": "success", "config": {
        "user_name": config.get("user_name", "אבא"),
        "version": config.get("version", "0.0.13"),
        "document_path": doc_path,
        "story_dir": story_dir,
        "active_story": active_story,
        "has_api_key": bool(config.get("gemini_api_key") and config.get("gemini_api_key") != "YOUR_GEMINI_API_KEY")
    }}

@app.get("/api/stories")
def get_stories():
    global config
    config = load_config()
    doc_path = config.get("document_path", "C:/Kolam/my_biography.docx")
    story_dir = os.path.dirname(doc_path)
    
    if not os.path.exists(story_dir):
        os.makedirs(story_dir, exist_ok=True)
        
    try:
        files = [f for f in os.listdir(story_dir) if f.endswith(".docx")]
        active_story = os.path.basename(doc_path)
        if active_story not in files:
            files.append(active_story)
        return {"status": "success", "stories": files, "active_story": active_story}
    except Exception as e:
        print(f"Error listing stories: {e}")
        return {"status": "error", "message": str(e), "stories": []}

@app.post("/api/select-story")
async def select_story(data: dict):
    global config
    config = load_config()
    
    story_name = data.get("story_name", "").strip()
    if not story_name:
        raise HTTPException(status_code=400, detail="שם הסיפור לא יכול להיות ריק")
        
    if not story_name.lower().endswith(".docx"):
        story_name += ".docx"
        
    doc_path = config.get("document_path", "C:/Kolam/my_biography.docx")
    story_dir = os.path.dirname(doc_path)
    
    new_path = os.path.join(story_dir, story_name).replace("\\", "/")
    config["document_path"] = new_path
    
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, ensure_ascii=False)
        
    print(f"Switched active story to: {new_path}")
    return {"status": "success", "active_story": story_name, "document_path": new_path}

@app.post("/api/open-story")
async def open_story():
    global config
    config = load_config()
    doc_path = config.get("document_path", "C:/Kolam/my_biography.docx")
    
    # Create it first if it doesn't exist
    if not os.path.exists(doc_path):
        try:
            os.makedirs(os.path.dirname(doc_path), exist_ok=True)
            from docx import Document
            doc = Document()
            title_p = doc.add_paragraph()
            title_p.paragraph_format.right_to_left = True
            run = title_p.add_run(f"הביוגרפיה של {config.get('user_name', 'אבא')}")
            run.bold = True
            doc.save(doc_path)
            print(f"Created new blank story document: {doc_path}")
        except Exception as e:
            return {"status": "error", "message": f"לא ניתן ליצור את הקובץ: {str(e)}"}
            
    try:
        os.startfile(os.path.abspath(doc_path))
        print(f"Opening story file in Word: {doc_path}")
        return {"status": "success", "message": "הקובץ נפתח ב-Word"}
    except Exception as e:
        print(f"Failed to open file in Word: {e}")
        return {"status": "error", "message": str(e)}

# Mount static files
os.makedirs("static", exist_ok=True)
app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    # Run server locally on port 8000
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
